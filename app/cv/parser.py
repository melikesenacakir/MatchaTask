from typing import Optional
from PyPDF2 import PdfReader
from docx import Document
import pdfplumber
import os


def parse_pdf(file_path: str) -> Optional[str]:

    if not os.path.exists(file_path): # control the file exists or not
        return None
    
    """   
    # Extraction with PdfReaderMethod
    try:
        reader = PdfReader(file_path) # read the file
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n" # extract the text from the page
        return text.strip() # clear and return the text
    except Exception as e:
        print(f"Error: {e}")
        return None
    """

    try:
        with pdfplumber.open(file_path) as pdf:
            text = ""
            for page in pdf.pages:
                text += page.extract_text() + "\n" # extract the text from the page
            return text.strip() # clear and return the text  
    except Exception as e:
        print(f"Error: {e}")
        return None   

  

def parse_docx(file_path: str) -> Optional[str]:

    if not os.path.exists(file_path):
        return None
    
    try:
        doc = Document(file_path)
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
    
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append("|".join(row_text))
        return "\n".join(text_parts).strip()
    except Exception as e:
        print(f"Error:{e}")
        return None




def parse_cv(file_path: str) -> Optional[str]:

    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return None
    
    file_ext = os.path.splitext(file_path)[1].lower()
    if file_ext == ".pdf":
        return parse_pdf(file_path)
    elif file_ext == ".docx":
        return parse_docx(file_path)
    else:
        print(f"Unsupported file format: {file_ext}")
        return None

if __name__ == "__main__":
    # PDF testi
    result_pdf = parse_cv("/Users/betulguner/Downloads/ResumeDataset/data/data/ENGINEERING/77828437.pdf")
    if result_pdf:
        print("PDF is parsed successfully!")
        print(result_pdf[:200])  # first 200 characters
    else:
        print("PDF can't be parsed!")
    
    print("\n" + "="*50 + "\n")
    
    # DOCX testi
    result_docx = parse_cv("/Users/betulguner/Downloads/exDocxCV.docx")
    if result_docx:
        print("DOCX is parsed successfully!")
        print(result_docx[:200])  # first 200 characters
    else:
        print("DOCX can't be parsed!")